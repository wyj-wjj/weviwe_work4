from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Literal

from app.core.config import Settings
from app.core.errors import AppError


ParseChoice = Literal["local", "ocr"]


@dataclass
class ExtractedBlock:
    type: str
    text: str


@dataclass
class ExtractedPage:
    page: int
    local_text: str = ""
    ocr_text: str = ""
    chosen: ParseChoice = "local"
    local_score: int = 0
    ocr_score: int = 0
    warning: str | None = None


@dataclass
class ExtractedDocument:
    raw_text: str
    blocks: list[ExtractedBlock] = field(default_factory=list)
    pages: list[ExtractedPage] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    parse_method: str = "unknown"
    parse_trace: dict[str, object] = field(default_factory=dict)


CONTENT_KEYWORDS = {
    "must_read": ["更新", "调整", "变更", "政策", "注意", "要求"],
    "base_script": ["产品", "流程", "客户", "口径", "说明"],
    "standard_script": ["场景", "推荐说法", "标准话术", "禁用说法", "注意事项"],
}

COMMON_PUNCTUATION = set("，。！？；：、,.!?;:()（）[]【】<>《》+-/%")


def normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    compact: list[str] = []
    previous_blank = False
    for line in lines:
        if not line:
            if not previous_blank:
                compact.append("")
            previous_blank = True
            continue
        compact.append(line)
        previous_blank = False
    return "\n".join(compact).strip()


def text_quality_score(text: str, *, content_type: str) -> int:
    stripped = text.strip()
    if not stripped:
        return 0
    readable_chars = [
        char
        for char in stripped
        if "\u4e00" <= char <= "\u9fff" or char.isascii() and char.isalnum() or char in COMMON_PUNCTUATION
    ]
    readable_ratio = len(readable_chars) / max(1, len(stripped))
    chinese_or_alnum_count = sum(1 for char in stripped if "\u4e00" <= char <= "\u9fff" or char.isalnum())
    score = min(45, chinese_or_alnum_count // 6)
    score += int(readable_ratio * 25)
    if any(keyword in stripped for keyword in CONTENT_KEYWORDS.get(content_type, [])):
        score += 12
    if "\n" in stripped and len([line for line in stripped.splitlines() if len(line.strip()) > 8]) >= 2:
        score += 8
    if "�" in stripped:
        score -= 25
    abnormal_count = sum(1 for char in stripped if not char.isspace() and char not in readable_chars)
    if abnormal_count:
        score -= min(25, abnormal_count * 2)
    single_char_lines = [line for line in stripped.splitlines() if len(line.strip()) == 1]
    if len(single_char_lines) >= 4:
        score -= 18
    if chinese_or_alnum_count < 20:
        score -= 12
    return max(0, min(100, score))


def choose_page_text(local_text: str, ocr_text: str, *, content_type: str, force_ocr: bool = False) -> ExtractedPage:
    local_score = text_quality_score(local_text, content_type=content_type)
    ocr_score = text_quality_score(ocr_text, content_type=content_type)
    chosen: ParseChoice
    warning = None
    if local_score >= ocr_score + 8:
        chosen = "local"
        if force_ocr and ocr_text.strip():
            warning = "已执行强制 OCR，本页最终采用本地文本，因为文本质量更高。"
    elif ocr_score >= local_score + 8:
        chosen = "ocr"
        warning = "本页本地文本疑似不完整，已采用 OCR。"
    else:
        chosen = "local" if len(local_text.strip()) >= len(ocr_text.strip()) else "ocr"
        if local_text.strip() and ocr_text.strip() and _text_difference_ratio(local_text, ocr_text) > 0.55:
            warning = "本页本地文本与 OCR 差异较大，请核对。"
    return ExtractedPage(
        page=0,
        local_text=local_text,
        ocr_text=ocr_text,
        chosen=chosen,
        local_score=local_score,
        ocr_score=ocr_score,
        warning=warning,
    )


def _text_difference_ratio(left: str, right: str) -> float:
    left_set = set(re.findall(r"[\w\u4e00-\u9fff]+", left))
    right_set = set(re.findall(r"[\w\u4e00-\u9fff]+", right))
    if not left_set and not right_set:
        return 0.0
    intersection = len(left_set & right_set)
    union = len(left_set | right_set)
    return 1 - intersection / max(1, union)


def extract_docx(
    file_bytes: bytes,
    *,
    content_type: str,
    dashscope_client,
    force_ocr: bool = False,
) -> ExtractedDocument:
    try:
        from docx import Document
    except ImportError as exc:
        raise AppError(code="document_parser_missing", message="DOCX parser is not installed.", status_code=503) from exc

    document = Document(io.BytesIO(file_bytes))
    blocks: list[ExtractedBlock] = []
    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if text:
            blocks.append(ExtractedBlock(type="paragraph", text=text))
    for table in document.tables:
        table_lines = []
        for row in table.rows:
            cells = [normalize_text(cell.text).replace("\n", " ") for cell in row.cells]
            line = " | ".join(cell for cell in cells if cell)
            if line:
                table_lines.append(line)
        if table_lines:
            blocks.append(ExtractedBlock(type="table", text="\n".join(table_lines)))

    warnings: list[str] = []
    ocr_used = False
    image_count = 0
    ocr_image_count = 0
    ocr_failed_count = 0
    local_block_count = len(blocks)
    skipped_image_ocr_count = 0
    for relationship in document.part.rels.values():
        if "image" not in relationship.reltype:
            continue
        image_count += 1
        image_bytes = relationship.target_part.blob
        if len(image_bytes) < 1024 and not force_ocr:
            continue
        try:
            ocr_text = normalize_text(
                dashscope_client.ocr_image(image_bytes=image_bytes, mime_type=_image_mime_type(image_bytes))
            )
        except Exception:
            ocr_failed_count += 1
            warnings.append("图片 OCR 失败，已保留本地文本；请人工核对图片中的内容。")
            continue
        if ocr_text:
            blocks.append(ExtractedBlock(type="image_ocr", text=ocr_text))
            ocr_used = True
            ocr_image_count += 1
    if skipped_image_ocr_count:
        warnings.append("文档包含图片，快速解析未执行图片 OCR；如图片含正文，请勾选强制 OCR 后重试。")
    if ocr_used:
        warnings.append("文档包含图片 OCR 内容，请核对识别结果。")

    raw_text = normalize_text("\n\n".join(block.text for block in blocks))
    if not raw_text:
        raise AppError(code="empty_document", message="未识别到有效文本，请检查文件。", status_code=422)
    parse_method = "docx_local"
    if ocr_used:
        parse_method = "docx_hybrid_ocr" if force_ocr else "docx_local_image_ocr"
    return ExtractedDocument(
        raw_text=raw_text,
        blocks=blocks,
        warnings=warnings,
        parse_method=parse_method,
        parse_trace={
            "file_type": "docx",
            "parse_method": parse_method,
            "local_block_count": local_block_count,
            "image_count": image_count,
            "ocr_image_count": ocr_image_count,
            "ocr_failed_count": ocr_failed_count,
            "ocr_page_count": 0,
        },
    )


def extract_pdf(
    file_bytes: bytes,
    *,
    content_type: str,
    parse_mode: str,
    force_ocr: bool,
    dashscope_client,
    settings: Settings,
) -> ExtractedDocument:
    try:
        import fitz
    except ImportError as exc:
        raise AppError(code="document_parser_missing", message="PDF parser is not installed.", status_code=503) from exc

    pdf = fitz.open(stream=file_bytes, filetype="pdf")
    if pdf.page_count > settings.content_import_max_pdf_pages:
        raise AppError(code="pdf_too_many_pages", message="PDF 页数过多，请拆分后上传。", status_code=422)

    local_texts = [normalize_text(page.get_text("text")) for page in pdf]
    ocr_page_indexes = _ocr_page_indexes(local_texts, content_type=content_type, parse_mode=parse_mode, force_ocr=force_ocr)
    if len(ocr_page_indexes) > settings.content_import_max_ocr_pages:
        raise AppError(code="ocr_page_limit_exceeded", message="OCR 页数过多，请拆分文件或选择快速解析。", status_code=422)

    pages: list[ExtractedPage] = []
    warnings: list[str] = []
    ocr_page_count = 0
    ocr_failed_count = 0
    for index, local_text in enumerate(local_texts):
        ocr_text = ""
        if index in ocr_page_indexes:
            page = pdf.load_page(index)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            try:
                ocr_text = normalize_text(
                    dashscope_client.ocr_image(image_bytes=pixmap.tobytes("png"), mime_type="image/png")
                )
                ocr_page_count += 1
            except Exception:
                ocr_failed_count += 1
                warnings.append(f"Page {index + 1} OCR failed; local text was kept.")
        page_choice = choose_page_text(local_text, ocr_text, content_type=content_type, force_ocr=force_ocr)
        page_choice.page = index + 1
        if page_choice.warning:
            warnings.append(f"第 {page_choice.page} 页{page_choice.warning}")
        pages.append(page_choice)

    chosen_texts = [page.local_text if page.chosen == "local" else page.ocr_text for page in pages]
    raw_text = normalize_text("\n\n".join(text for text in chosen_texts if text.strip()))
    if not raw_text:
        raise AppError(code="empty_document", message="未识别到有效文本，请检查文件。", status_code=422)
    parse_method = "pdf_hybrid_enhanced" if parse_mode == "enhanced" or force_ocr else "pdf_hybrid_fast"
    return ExtractedDocument(
        raw_text=raw_text,
        blocks=[ExtractedBlock(type="page", text=text) for text in chosen_texts if text.strip()],
        pages=pages,
        warnings=warnings,
        parse_method=parse_method,
        parse_trace={
            "file_type": "pdf",
            "parse_method": parse_method,
            "local_block_count": len([text for text in local_texts if text.strip()]),
            "image_count": 0,
            "ocr_image_count": 0,
            "ocr_failed_count": ocr_failed_count,
            "ocr_page_count": ocr_page_count,
        },
    )


def _ocr_page_indexes(
    local_texts: list[str],
    *,
    content_type: str,
    parse_mode: str,
    force_ocr: bool,
) -> set[int]:
    if force_ocr or parse_mode == "enhanced":
        return set(range(len(local_texts)))
    scored = [
        (index, text_quality_score(text, content_type=content_type), text)
        for index, text in enumerate(local_texts)
    ]
    selected = {0} if local_texts else set()
    selected.update(index for index, score, _text in scored if score < 35)
    selected.update(index for index, _score, text in scored if "|" in text or "\t" in text)
    for index, _score, _text in sorted(scored, key=lambda item: item[1])[:3]:
        selected.add(index)
    return selected


def _image_mime_type(image_bytes: bytes) -> str:
    if image_bytes.startswith(b"\xff\xd8"):
        return "image/jpeg"
    if image_bytes.startswith(b"\x89PNG"):
        return "image/png"
    return "application/octet-stream"
